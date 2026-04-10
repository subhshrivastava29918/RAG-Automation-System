"""
Document Ingestion Pipeline
Supports: .txt, .md, .pdf, .docx, URLs, raw strings
"""

import os
import re
import hashlib
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from urllib.parse import urlparse

from .engine import Document

logger = logging.getLogger(__name__)


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_md(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    # Strip markdown syntax for cleaner text
    content = re.sub(r'#{1,6}\s', '', content)
    content = re.sub(r'\*{1,2}(.+?)\*{1,2}', r'\1', content)
    content = re.sub(r'`{1,3}.*?`{1,3}', '', content, flags=re.DOTALL)
    return content


def _extract_pdf(path: str) -> str:
    try:
        import PyPDF2
        text = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"[Page {page_num + 1}]\n{page_text}")
        return "\n\n".join(text)
    except ImportError:
        logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction error for {path}: {e}")
        return ""


def _extract_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error for {path}: {e}")
        return ""


def _extract_url(url: str) -> str:
    try:
        import requests
        from bs4 import BeautifulSoup
        headers = {"User-Agent": "Mozilla/5.0 (RAG-Bot/1.0)"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        # Remove scripts, styles, nav, footer
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        return "\n".join(lines)
    except ImportError:
        logger.error("requests/beautifulsoup4 not installed.")
        return ""
    except Exception as e:
        logger.error(f"URL extraction error for {url}: {e}")
        return ""


class DocumentIngester:
    """
    Ingest documents from various sources and produce Document objects.
    Supports deduplication via content hashing.
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".json"}

    def __init__(self):
        self._seen_hashes: set = set()

    def _make_doc_id(self, content: str, source: str) -> str:
        return hashlib.sha256(f"{source}:{content[:200]}".encode()).hexdigest()[:16]

    def _is_duplicate(self, content: str) -> bool:
        h = hashlib.md5(content.encode()).hexdigest()
        if h in self._seen_hashes:
            return True
        self._seen_hashes.add(h)
        return False

    def ingest_text(self, text: str, source: str = "raw", metadata: Optional[Dict] = None) -> Optional[Document]:
        """Ingest raw text string."""
        text = text.strip()
        if not text:
            return None
        if self._is_duplicate(text):
            logger.warning(f"Duplicate content from source={source}, skipping.")
            return None
        doc_id = self._make_doc_id(text, source)
        return Document(
            doc_id=doc_id,
            content=text,
            source=source,
            doc_type="text",
            metadata=metadata or {},
        )

    def ingest_file(self, path: str, metadata: Optional[Dict] = None) -> Optional[Document]:
        """Ingest a local file (txt, md, pdf, docx)."""
        p = Path(path)
        if not p.exists():
            logger.error(f"File not found: {path}")
            return None

        ext = p.suffix.lower()
        extractors = {
            ".txt": _extract_txt,
            ".md": _extract_md,
            ".pdf": _extract_pdf,
            ".docx": _extract_docx,
            ".doc": _extract_docx,
        }

        extractor = extractors.get(ext)
        if extractor is None:
            # Try reading as plain text for unknown extensions
            try:
                content = _extract_txt(path)
            except Exception:
                logger.error(f"Unsupported file type: {ext}")
                return None
        else:
            content = extractor(path)

        if not content or not content.strip():
            logger.warning(f"No content extracted from {path}")
            return None

        meta = {"filename": p.name, "extension": ext, **(metadata or {})}
        if self._is_duplicate(content):
            logger.warning(f"Duplicate file content: {path}, skipping.")
            return None

        doc_id = self._make_doc_id(content, str(p))
        return Document(
            doc_id=doc_id,
            content=content,
            source=str(p.resolve()),
            doc_type=ext.lstrip("."),
            metadata=meta,
        )

    def ingest_directory(
        self,
        dir_path: str,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
        metadata: Optional[Dict] = None,
    ) -> List[Document]:
        """Ingest all supported files from a directory."""
        exts = {e.lower() for e in (extensions or self.SUPPORTED_EXTENSIONS)}
        base = Path(dir_path)
        pattern = "**/*" if recursive else "*"
        documents = []
        for file_path in base.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in exts:
                doc = self.ingest_file(str(file_path), metadata=metadata)
                if doc:
                    documents.append(doc)
                    logger.info(f"Ingested: {file_path.name}")
        return documents

    def ingest_url(self, url: str, metadata: Optional[Dict] = None) -> Optional[Document]:
        """Ingest content from a web URL."""
        parsed = urlparse(url)
        if not parsed.scheme.startswith("http"):
            logger.error(f"Invalid URL: {url}")
            return None

        content = _extract_url(url)
        if not content.strip():
            return None

        meta = {"url": url, "domain": parsed.netloc, **(metadata or {})}
        if self._is_duplicate(content):
            logger.warning(f"Duplicate URL content: {url}")
            return None

        doc_id = self._make_doc_id(content, url)
        return Document(
            doc_id=doc_id,
            content=content,
            source=url,
            doc_type="web",
            metadata=meta,
        )

    def ingest_json_list(self, data: List[Dict], text_field: str = "text", metadata_fields: Optional[List[str]] = None) -> List[Document]:
        """Ingest a list of dicts with a text field."""
        documents = []
        for item in data:
            text = item.get(text_field, "")
            if not text:
                continue
            meta = {}
            if metadata_fields:
                meta = {k: item.get(k, "") for k in metadata_fields}
            source = item.get("source", item.get("id", "json"))
            doc = self.ingest_text(text, source=str(source), metadata=meta)
            if doc:
                documents.append(doc)
        return documents
